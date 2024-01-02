import time

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import json
import os
import tkinter as tk
import customtkinter as CTk
from PIL import Image
import datetime
from collections import Counter
from docx import Document
from docx.shared import Inches
import comtypes.client
from functools import reduce


class WorkersList(CTk.CTkScrollableFrame):
    def __init__(self, master, **kwargs):
        super().__init__(master, **kwargs)

        # add widgets onto the frame...
        self.search_frame = CTk.CTkFrame(master=self, width=750, height=95,
                                         fg_color="transparent", border_width=0, border_color='black')
        self.search_frame.grid(row=0, column=0, padx=(0, 0), pady=(0, 0), sticky="nsew")

        self.list_frame = CTk.CTkFrame(master=self, width=750, height=500,
                                       fg_color="transparent")
        self.list_frame.grid(row=1, column=0, padx=(0, 0), pady=(0, 0), sticky="nsew")

        self.left_column = CTk.CTkFrame(master=self.list_frame, width=100, height=500,
                                        fg_color="transparent")
        self.left_column.grid(row=0, column=0, padx=(0, 0), pady=(0, 0), sticky="nsew")

        self.right_column = CTk.CTkFrame(master=self.list_frame, width=370, height=500,
                                         fg_color="red")
        #self.right_column.grid(row=0, column=1, padx=(0, 0), pady=(0, 0), sticky="nsew")

        self.tab_name = CTk.CTkLabel(self.search_frame, text="Выберите до 4 сотрудников чтобы увидеть их статистику",
                                     fg_color="transparent", justify='left',
                                     width=375, height=34, anchor="w", font=("Arial", 15))
        self.tab_name.grid(row=0, column=0, padx=(5, 0), pady=(3, 3), sticky="nsew")

        self.search_entry = CTk.CTkEntry(self.search_frame, placeholder_text="Посик", width=277,
                                         height=34, corner_radius=20)
        self.search_entry.grid(row=0, column=1, padx=(52, 3), pady=(3, 3), sticky="nsew")

        with open('data/available_workers.json', 'r', encoding='utf8') as workers_file:
            self.available_workers = json.load(workers_file)

        self.workers_to_show = list(self.available_workers.values())[0].copy()
        self.workers_checkboxes = list()
        self.workers_with_stat = list()

        def workers_search():
            request = self.search_entry.get()
            if request == '':
                self.workers_to_show = list(self.available_workers.values())[0].copy()
                self.clear_workers_list()
                self.draw_workers_list()
            else:
                self.clear_workers_list()
                self.workers_to_show = list()
                for worker in list(self.available_workers.values())[0]:
                    if request in worker:
                        self.workers_to_show.append(worker)
                self.draw_workers_list()

        self.search_button = CTk.CTkButton(self.search_frame, text="Найти", command=workers_search,
                                           width=40, height=24, corner_radius=17)
        self.search_button.grid(row=0, column=3, padx=(5, 3), pady=(8, 8), sticky="nsew")

        self.workers_box_state = dict()

        self.draw_workers_list()

    def clear_workers_list(self):
        for worker in self.workers_checkboxes:
            worker.grid_forget()
        self.workers_checkboxes = list()

    def draw_workers_list(self):
        cur_col = 0
        cur_row = 0
        for ind, worker in enumerate(self.workers_to_show):
            def checkbox_event_wrapper(cur_worker: str):
                def checkbox_event():
                    if self.workers_box_state[cur_worker].get() == 'on':
                        if len(self.workers_with_stat) <= 10:
                            self.workers_with_stat.append(cur_worker)
                    else:
                        self.workers_with_stat.remove(cur_worker)
                    print(self.workers_with_stat)
                return checkbox_event
            if self.workers_box_state.get(worker) is None:
                self.workers_box_state[worker] = CTk.StringVar(value="off")
            else:
                pass

            checkbox = CTk.CTkCheckBox(self.left_column, text=worker,
                                       command=checkbox_event_wrapper(worker),
                                       variable=self.workers_box_state[worker], onvalue="on", offvalue="off",
                                       width=100, height=44, border_width=1.5, corner_radius=0)
            checkbox.grid(row=ind, column=0, padx=(5, 0), pady=(5, 0), sticky="w")

            self.workers_checkboxes.append(checkbox)


class MyTabView(CTk.CTkTabview):
    def __init__(self, master, **kwargs):
        super().__init__(master, **kwargs)

        # create tabs
        self.add("Быстрый рассчет")
        self.add("Выбор сотрудников")
        self.add("Статистика сотрудников")
        self.add("Настройки границ оценок")
        self.add("Настройки алгоритма оценки")

        # add widgets on tabs
        self.directory = os.getcwd()
        self.ready_to_calc = False
        self.path_to_data = None
        self.now_time = str(datetime.datetime.now()).split()[1].split('.')[0]
        self.should_extend = False
        self.should_update = True

        self.active_features_list = list()
        self.active_regimes = list()
        self.tables = list()
        self.distribution_cells = list()

        self.weight_dict = dict()
        self.available_workers = dict()
        self.features = {
            "Грамматика": 0.1,
            "Лексика": 0.1,
            "Контекст": 0.1,
            "Логотип": 0.1,
            "Внешний отправитель": 0.1,
            "Давление": 0.1,
            "Структура": 0.1,
            "Гиперссылка": 0.1,
        }

        self.features_state = {
            "Грамматика": "off",
            "Лексика": "off",
            "Контекст": "off",
            "Логотип": "off",
            "Внешний отправитель": "off",
            "Давление": "off",
            "Структура": "off",
            "Гиперссылка": "off",
        }

        self.regime_state = {
            "Вычиление оценок": "off",
            "Оценка результатов": "off",
        }

        self.settings_paths = {
            "settings": '',
            "workers": '',
            "weights": '',
        }

        self.data_for_calculation = {
            'Ends': '',
            'Already_calc': [],
        }

        for file in os.listdir(self.directory + '\\data\\'):
            filename = os.fsdecode(file)
            if filename == "settings.json":
                self.settings_paths['settings'] = f"{self.directory}/data/{filename}"
            elif filename == "available_workers.json":
                self.settings_paths['workers'] = f"{self.directory}/data/{filename}"
            elif filename == "weight_dict.json":
                self.settings_paths['weights'] = f"{self.directory}/data/{filename}"
            else:
                pass

        with open(self.settings_paths["weights"], 'r', encoding='utf8') as weight_file:
            self.weight_dict = json.load(weight_file)

        self.settings_paths_enter = {
            "settings": '',
            "workers": '',
            "weights": '',
        }

        self.paths_name = ['настроек', 'сотрудников', 'весов']

        # Добавление фреймов Быстрый рассчет
        self.top_frame = CTk.CTkFrame(master=self.tab("Быстрый рассчет"), width=690, height=95, fg_color="transparent")
        self.top_frame.grid(row=0, column=0, padx=(5, 5), pady=(5, 0), sticky="nsew")

        self.mid_frame = CTk.CTkFrame(master=self.tab("Быстрый рассчет"), width=690, height=395, fg_color="transparent")
        self.mid_frame.grid(row=1, column=0, padx=(5, 5), pady=(0, 5), sticky="nsew")

        self.feats_frame = CTk.CTkFrame(master=self.mid_frame, width=230, height=395, fg_color="transparent")
        self.feats_frame.grid(row=0, column=0, padx=(0, 0), pady=(0, 0), sticky="nsew")

        self.mid_right_frame = CTk.CTkFrame(master=self.mid_frame, width=460, height=395, fg_color="transparent")
        self.mid_right_frame.grid(row=0, column=1, padx=(0, 5), pady=(0, 0), sticky="nsew")

        self.regime_frame = CTk.CTkFrame(master=self.mid_right_frame, width=230, height=197.5, fg_color="transparent")
        self.regime_frame.grid(row=0, column=0, padx=(0, 0), pady=(0, 0), sticky="nsew")

        self.grade_frame = CTk.CTkFrame(master=self.mid_right_frame, width=230, height=197.5, fg_color="transparent")
        self.grade_frame.grid(row=0, column=1, padx=(0, 0), pady=(0, 0), sticky="nsew")

        self.info_frame = CTk.CTkFrame(master=self.mid_right_frame, width=510, height=225, fg_color="transparent")
        self.info_frame.grid(row=1, column=0, columnspan=2, padx=(0, 2.5), pady=(20, 0), sticky="nsew")

        # Добавление фреймов Информация о сотрудниках
        self.scroll_info_frame = WorkersList(master=self.tab("Выбор сотрудников"), width=815, height=530)
        self.scroll_info_frame.grid(row=0, column=0, padx=0, pady=0)

        # Добавление фреймов Статистика сотрудников
        self.features_label = CTk.CTkLabel(self.tab("Статистика сотрудников"), text="Выберите хотя бы одного "
                                                                                    "сотрудника в разделе "
                                                                                    "'Выбор сотрудников' и нажмите "
                                                                                    "Обновить",
                                           fg_color="transparent",
                                           width=225, height=44, anchor="w", font=("Arial", 15))
        self.features_label.grid(row=0, column=0, padx=(5, 0), pady=(5, 0), sticky="nsew")

        self.stat_frame = CTk.CTkFrame(master=self.tab("Статистика сотрудников"), width=815, height=480,
                                       fg_color="transparent")
        self.stat_frame.grid(row=2, column=0, padx=(0, 0), pady=(5, 0), sticky="nsew")

        self.cells_list = list()

        def update_command():
            self.update_cells()
            if self.scroll_info_frame.workers_with_stat:
                for i in range(4):
                    try:
                        worker = self.scroll_info_frame.workers_with_stat[i]
                    except IndexError:
                        worker = ''
                    print(worker)
                    cur_cell_dict = self.cells_list[i]
                    if len(worker) > 25:
                        font_set = ("Arial", 10)
                    else:
                        font_set = ("Arial", 15)
                    name_label = CTk.CTkLabel(cur_cell_dict['name_frame'], text=worker,
                                              fg_color="transparent",
                                              width=410, height=38, font=font_set)
                    name_label.grid(row=0, column=0, padx=(3, 0), pady=(0, 0), sticky="nsew")

                    grade_label = CTk.CTkLabel(cur_cell_dict['diff_val_frame'], text='Оценка',
                                               fg_color="transparent",
                                               width=204, height=38, font=font_set)
                    grade_label.grid(row=0, column=0, padx=(3, 0), pady=(0, 0), sticky="nsew")

                    distr_name_label = CTk.CTkLabel(cur_cell_dict['distr_name_frame'], text='Распределение',
                                                    fg_color="transparent",
                                                    width=204, height=38, font=font_set)
                    distr_name_label.grid(row=0, column=0, padx=(3, 0), pady=(0, 0), sticky="nsew")

                    distr_val_label = CTk.CTkLabel(cur_cell_dict['distr_val_frame'], text='Значения',
                                                   fg_color="transparent",
                                                   width=204, height=38, font=font_set)
                    distr_val_label.grid(row=0, column=0, padx=(3, 0), pady=(0, 0), sticky="nsew")

                    pychart_label = CTk.CTkLabel(cur_cell_dict['pychart_frame'], text='PyCahrt',
                                                 fg_color="transparent",
                                                 width=204, height=191, font=font_set)
                    pychart_label.grid(row=0, column=0, padx=(3, 0), pady=(0, 0), sticky="nsew")

        self.update_frame = CTk.CTkFrame(master=self.tab("Статистика сотрудников"), width=100, height=44,
                                         fg_color="transparent")
        self.update_frame.grid(row=1, column=0, padx=(5, 0), pady=(0, 0), sticky="nsew")

        self.update_button = CTk.CTkButton(master=self.update_frame, width=100, height=32,
                                           command=update_command, corner_radius=16, text="Обновить")
        self.update_button.grid(row=0, column=0, padx=(5, 0), pady=(0, 0), sticky="nsew")

        # Добавление фреймов Найтсройки 1
        self.settings_test_label = CTk.CTkLabel(self.tab("Настройки границ оценок"), text="В процессе разработки",
                                                fg_color="transparent",
                                                width=225, height=44, anchor="w", font=("Arial", 15))
        self.settings_test_label.grid(row=0, column=0, padx=(5, 0), pady=(5, 0), sticky="nsew")

        # Добавление фреймов Найтсройки 2
        self.settings_test_label = CTk.CTkLabel(self.tab("Настройки алгоритма оценки"), text="В процессе разработки",
                                                fg_color="transparent",
                                                width=225, height=44, anchor="w", font=("Arial", 15))
        self.settings_test_label.grid(row=0, column=0, padx=(5, 0), pady=(5, 0), sticky="nsew")

        # Добавление интеррактивных элементов
        self.features_label = CTk.CTkLabel(self.feats_frame, text="Признаки", fg_color="transparent",
                                           width=225, height=44, anchor="w", font=("Arial", 15))
        self.features_label.grid(row=0, column=0, padx=(5, 0), pady=(5, 5), sticky="nsew")
        for ind, key in enumerate(self.features):
            def checkbox_event_wrapper(feature: str):
                def checkbox_event():
                    if self.features_state[feature].get() == 'on':
                        self.active_features_list.append(feature)
                    else:
                        self.active_features_list.remove(feature)
                    self.change_pred_difficult()
                    self.change_pred_distribution()
                return checkbox_event

            self.features_state[key] = CTk.StringVar(value="off")
            self.checkbox = CTk.CTkCheckBox(self.feats_frame, text=key, command=checkbox_event_wrapper(key),
                                            variable=self.features_state[key], onvalue="on", offvalue="off",
                                            width=225, height=44, border_width=1.5, corner_radius=0)
            self.checkbox.grid(row=ind + 1, column=0, padx=(5, 0), pady=(5, 0), sticky="nsew")

        self.regime_label = CTk.CTkLabel(self.regime_frame, text="Режим работы", fg_color="transparent",
                                         width=225, height=44, anchor="w", font=("Arial", 15))
        self.regime_label.grid(row=0, column=0, padx=(5, 0), pady=(5, 0), sticky="nsew")
        for ind, key in enumerate(self.regime_state):
            def checkbox_event_regime_wrapper(regime: str):
                def checkbox_event():
                    if self.regime_state[regime].get() == 'on':
                        self.active_regimes.append(regime)
                    else:
                        self.active_regimes.remove(regime)
                return checkbox_event

            self.regime_state[key] = CTk.StringVar(value="off")
            self.checkbox = CTk.CTkCheckBox(self.regime_frame, text=key, command=checkbox_event_regime_wrapper(key),
                                            variable=self.regime_state[key], onvalue="on", offvalue="off",
                                            width=170, height=44, border_width=1.5, corner_radius=0)
            self.checkbox.grid(row=ind + 1, column=0, padx=(5, 0), pady=(5, 0), sticky="nsew")

        self.info_text = CTk.CTkTextbox(master=self.info_frame,
                                        width=590, height=185, corner_radius=0, state='disabled')
        self.info_text.grid(row=0, column=0, padx=(5, 5), pady=(5, 0), sticky="nsew")

        for ind, key in enumerate(self.settings_paths):
            data = self.settings_paths[key]
            settings_path = CTk.CTkEntry(self.grade_frame, placeholder_text=data if data else "Файл не найден",
                                         width=250, height=31)

            self.settings_paths_enter[key] = settings_path

            def get_btn_event_wrapper(path: str):
                def get_paths_button_event():
                    entered_data = self.settings_paths_enter[path].get()
                    if entered_data:
                        self.settings_paths[path] = entered_data

                return get_paths_button_event

            self.get_paths_btn = CTk.CTkButton(self.grade_frame, text="Обновить",
                                               command=get_btn_event_wrapper(key),
                                               width=77.5,
                                               height=31, corner_radius=16)

        self.grade_label = CTk.CTkLabel(self.grade_frame, text="Ожидаемая сложность: ", fg_color="transparent",
                                        width=225, height=44, anchor="w", font=("Arial", 15))
        self.grade_label.grid(row=0, column=0, padx=(0, 0), pady=(5, 0), sticky="nsew")

        self.grade_label_val = CTk.CTkLabel(self.grade_frame, text="10", fg_color="transparent",
                                            height=44, anchor="w", font=("Arial", 15))
        self.grade_label_val.grid(row=0, column=1, padx=(0, 0), pady=(5, 0), sticky="nsew")

        self.answers_label = CTk.CTkLabel(self.grade_frame, text="Ожидаемое распределение", fg_color="transparent",
                                          height=44, anchor="w", font=("Arial", 15))
        self.answers_label.grid(row=1, column=0, columnspan=2, padx=(0, 0), pady=(0, 0), sticky="nsew")

        self.answers_table = CTk.CTkFrame(master=self.grade_frame, width=225, height=88, fg_color="transparent")
        self.answers_table.grid(row=2, column=0, columnspan=2, padx=(0, 0), pady=(5, 0), sticky="nsew")

        distribution = self.calc_distribution()
        for i in range(6):
            table_cell_number = CTk.CTkLabel(self.answers_table, text=str(i+1), fg_color="transparent", width=39,
                                             height=42)
            table_cell_number.grid(row=0, column=i, padx=(0, 0), pady=(0, 0), sticky="nsew")

            table_cell_val = CTk.CTkLabel(self.answers_table, text="%.3f" % distribution[i], fg_color="transparent",
                                          width=39,
                                          height=43)
            table_cell_val.grid(row=1, column=i, padx=(0, 0), pady=(0, 0), sticky="nsew")
            self.distribution_cells.append(table_cell_val)

        self.data_path_entry = CTk.CTkEntry(self.top_frame, placeholder_text="Путь до данных", width=460, height=31)
        self.data_path_entry.grid(row=0, column=0, padx=(5, 0), pady=(31, 31), sticky="nsew")

        def get_path_button_event():
            entered_data = self.data_path_entry.get()
            if entered_data:
                self.log_message('Путь до данных обновлен')
                self.path_to_data = str(entered_data).replace('\\', '/')

        self.get_path_button = CTk.CTkButton(self.top_frame, text="Обновить", command=get_path_button_event, width=77.5,
                                             height=31, corner_radius=16)
        self.get_path_button.grid(row=0, column=1, padx=(20, 0), pady=(31, 31), sticky="nsew")

        def calc_btn_event():
            try:
                if self.active_regimes:
                    self.prepare_settings()
                else:
                    self.log_message('Выберете режим работы')
                    return
            except ValueError:
                self.log_message('Ошибка в пути до данных')
                return
            try:
                self.tables = self.prepare_data()
                self.log_message('Завершено')
            except FileNotFoundError:
                self.log_message('Ошибка в пути до данных')
                return
            res = self.calc_result(self.tables)
            print(res)
            try:
                self.work_with_data(res)
            except ValueError:
                self.log_message('Данные не найдены')

        self.calc_btn = CTk.CTkButton(self.info_frame, text="Вычислить",
                                      command=calc_btn_event,
                                      width=100,
                                      height=30, corner_radius=16)
        self.calc_btn.grid(row=1, column=0, padx=(0, 5), pady=(15, 5), sticky="e")

    @staticmethod
    def change_data(value: str):
        return str(1) if value.split()[0] == 'есть' else str(0)

    def calc_difficult(self):
        features_weight_sum = 0
        for feature in self.active_features_list:
            features_weight_sum += self.weight_dict[feature]
        return int(np.round((1 - features_weight_sum) * 10))

    def change_pred_difficult(self):
        self.grade_label_val.configure(text=str(self.calc_difficult()))

    def clear_cells(self):
        for cell in self.cells_list:
            for cell_data in cell.values():
                cell_data.grid_forget()

    def update_cells(self):
        self.clear_cells()
        self.cells_list = list()

        cur_row = 0
        cur_col = 0
        cur_len = len(self.scroll_info_frame.workers_with_stat)
        steps = cur_len if cur_len <= 4 else 4
        for i in range(steps):
            cell_frame = CTk.CTkFrame(master=self.stat_frame, width=405, height=230,
                                      fg_color="transparent")
            cell_frame.grid(row=cur_row, column=cur_col, padx=(3, 3), pady=(3, 3), sticky="nsew")

            name_frame = CTk.CTkFrame(master=cell_frame, width=405, height=38,
                                      fg_color="white")
            name_frame.grid(row=0, column=0, columnspan=2, padx=(0, 0), pady=(0, 0), sticky="nsew")

            diff_val_frame = CTk.CTkFrame(master=cell_frame, width=204, height=38,
                                          fg_color="white")
            diff_val_frame.grid(row=1, column=1, padx=(0, 0), pady=(0, 0), sticky="nsew")

            distr_name_frame = CTk.CTkFrame(master=cell_frame, width=204, height=38,
                                            fg_color="white")
            distr_name_frame.grid(row=2, column=1, padx=(0, 0), pady=(0, 0), sticky="nsew")

            distr_val_frame = CTk.CTkFrame(master=cell_frame, width=204, height=115,
                                           fg_color="white")
            distr_val_frame.grid(row=3, column=1, padx=(0, 0), pady=(0, 0), sticky="nsew")

            pychart_frame = CTk.CTkFrame(master=cell_frame, width=204, height=191,
                                         fg_color="white")
            pychart_frame.grid(row=1, column=0, rowspan=3, padx=(0, 0), pady=(0, 0), sticky="nsew")

            dict_to_append = {
                "cell_frame": cell_frame,
                "name_frame": name_frame,
                "diff_val_frame": diff_val_frame,
                "distr_name_frame": distr_name_frame,
                "distr_val_frame": distr_val_frame,
                "pychart_frame": pychart_frame,
            }

            self.cells_list.append(dict_to_append)

            cur_col += 1
            if cur_col > 1:
                cur_col = 0
                cur_row += 1

    def change_pred_distribution(self):
        distribution = self.calc_distribution()
        for i in range(6):
            self.distribution_cells[i].configure(text="%.3f" % distribution[i])

    def calc_distribution(self):
        distribution = [0.025, 0.01, 0.05, 0.015, 0.7, 0.2]
        difficult = self.calc_difficult()
        if difficult - 5 > 0:
            for i in range(abs(difficult-5)):
                if distribution[5] > 0.06:
                    distribution[4] -= 0.05
                    distribution[5] -= 0.05
                    distribution[0] += 0.025
                    distribution[1] += 0.025
                    distribution[2] += 0.025
                    distribution[3] += 0.025
                else:
                    distribution[4] -= 0.1
                    distribution[0] += 0.025
                    distribution[1] += 0.025
                    distribution[2] += 0.025
                    distribution[3] += 0.025
        elif difficult - 5 < 0:
            for i in range(abs(difficult-5)):
                if distribution[1] > 0.026:
                    distribution[4] += 0.05
                    distribution[5] += 0.05
                    distribution[0] -= 0.025
                    distribution[1] -= 0.025
                    distribution[2] -= 0.025
                    distribution[3] -= 0.025
                elif distribution[1] > 0.006:
                    distribution[5] += 0.02
                    distribution[1] -= 0.005
                    distribution[2] -= 0.015
                elif distribution[3] > 0.006:
                    distribution[4] += 0.01
                    distribution[5] += 0.02
                    distribution[0] -= 0.01
                    distribution[2] -= 0.01
                    distribution[3] -= 0.01
                elif distribution[0] > 0.006:
                    distribution[5] += 0.12
                    distribution[4] -= 0.1
                    distribution[0] -= 0.01
                    distribution[2] -= 0.01
                elif distribution[2] > 0.006:
                    distribution[5] += 0.105
                    distribution[4] -= 0.1
                    distribution[2] -= 0.005
        else:
            pass
        print(distribution)
        return distribution

    def log_message(self, text):
        self.info_text.configure(state='normal')
        self.now_time = str(datetime.datetime.now()).split()[1].split('.')[0]
        self.info_text.insert("end", f"[{self.now_time}]: {text}\n")
        self.info_text.configure(state='disabled')

    def prepare_settings(self):
        try:
            for path in self.settings_paths.values():
                assert path
                assert os.path.exists(path)
            print(self.path_to_data)
            assert self.path_to_data is not None
            assert os.path.exists(self.path_to_data)
        except AssertionError:
            # self.info_text.delete("0.0", "end")\
            self.now_time = str(datetime.datetime.now()).split()[1].split('.')[0]
            self.info_text.insert("end", f"[{self.now_time}]: Ошибка, укажите пути до всех файлов\n")
            raise ValueError

        with open(self.settings_paths["settings"], 'r', encoding='utf8') as settings_file:
            settings = json.load(settings_file)
            self.data_for_calculation['Ends'] = settings['Ends']
            self.data_for_calculation['Already_calc'] = settings['Already_calc']

        with open(self.settings_paths["weights"], 'r', encoding='utf8') as weight_file:
            self.weight_dict = json.load(weight_file)

        with open(self.settings_paths["workers"], 'r', encoding='utf8') as workers_file:
            self.available_workers = json.load(workers_file)

    def prepare_data(self) -> list:
        tables = list()
        main_columns = ['Переход_по_ссылке', 'Ввод_данных', 'Сообщение_в_ЦПКБ']
        for file in os.listdir(self.path_to_data):
            try:
                if file.split('.')[-1] in ['xlsx', 'csv']:
                    cur_table = pd.read_excel(self.path_to_data + file)
                    cur_table[main_columns] = cur_table[main_columns].fillna('нет')
                    if {'Должность', 'Переход_по_ссылке', 'Ввод_данных', 'Сообщение_в_ЦПКБ'} in set(cur_table.columns):
                        error_file = file.split('.')[-1]
                        raise TypeError
                    for column in main_columns:
                        if cur_table[column].dtypes is not int:
                            for col_value in cur_table[column].unique():
                                if col_value.split()[0] not in ['есть', 'нет']:
                                    raise TypeError
                    tables.append(cur_table)
                else:
                    raise ValueError
            except ValueError:
                print('Посторонние файлы в папке')
            except TypeError:
                print(f'Неверный формат данных')
        return tables

    def calc_result(self, tables: list) -> tuple:
        end_tables = list()
        end_dict = self.data_for_calculation['Ends']
        difficult = self.calc_difficult()
        main_columns = ['Переход_по_ссылке', 'Ввод_данных', 'Сообщение_в_ЦПКБ']
        for table in tables:
            for column in main_columns:
                table[column] = table[column].map(self.change_data)
            table['Исход'] = table['Переход_по_ссылке'] + table['Ввод_данных'] + table['Сообщение_в_ЦПКБ']
            table[f'{difficult}'] = table['Исход'].map(lambda value: end_dict[value])
            end_tables.append(table[['Должность', f'{difficult}']].groupby('Должность').mean())
        if difficult not in self.data_for_calculation['Already_calc']:
            self.data_for_calculation['Already_calc'].append(difficult)
            self.should_extend = True
            self.should_update = False
        return end_tables, difficult

    def make_grades_table(self, cur_table, res):
        general_table = pd.read_excel('general_table.xlsx')
        grades_table = pd.DataFrame.from_dict({"Должность": cur_table.index})
        difficult = res[1]
        try:
            grades_table['Результат'] = cur_table[f'{difficult}']
        except ValueError:
            print(cur_table[f'{difficult}'])
            grades_table['Результат'] = list(reduce(lambda x, y: x + y, cur_table[f'{difficult}']))
        grades = list()
        for cur_val, grade in zip(cur_table[f'{difficult}'],
                                  general_table[f'{difficult}']):
            grades.append('Успешно прошел') if cur_val >= grade else grades.append('Не прошел')
        try:
            grades_table['Оценка'] = grades
        except ValueError:
            self.remake_workers(cur_table)
            grades_table['Оценка'] = grades
        grades_table.to_excel('grades_table.xlsx')

    @staticmethod
    def show_top(cur_table):
        dirty_list = list(sorted(list(zip(cur_table.index, cur_table.values)), key=lambda val: val[1][0]))
        clear_list = list(filter(lambda val: val[1][0] > -1, dirty_list))[:5]
        print(clear_list)
        return clear_list

    @staticmethod
    def calc_mean(cur_table):
        with open('grouped_workers.json', 'r', encoding='utf8') as group_data_file:
            group_data = json.load(group_data_file)

        group_dict = {
            0: list(),
            1: list(),
            2: list(),
        }
        for key, val in zip(cur_table.index, cur_table.values):
            group = group_data[key]
            group_dict[group].append(val)
        return np.mean(group_dict[0]), np.mean(group_dict[1]), np.mean(group_dict[2])

    @staticmethod
    def make_compare_plot(x: dict, y: dict):
        plt.figure(figsize=(10, 6))
        plt.hist(y, bins=50, color='orange', density=True)
        plt.hist(x, bins=50, color='red', density=True)
        plt.title("Распределение ответов сотрудников")
        plt.xlabel("Ответы сотрудников")
        plt.savefig('plt_1.png')

    @staticmethod
    def remake_workers(cur_table):
        with open('data/available_workers.json', 'w', encoding='utf8') as json_file:
            json.dump({"Должность": list(cur_table.index)}, json_file, ensure_ascii=False)

    def make_document(self, general_text, top_text):
        document = Document()

        main_head = document.add_heading('Отчет об оценки результатов\n', 0)
        main_head.alignment = 1

        document.add_heading('Общая информация', 2)
        document.add_paragraph(general_text)

        document.add_heading('Топ 5 худших результатов', 2)
        document.add_paragraph(top_text)

        document.add_heading('Распределение ответов сотрудников', 2)
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture('plt_1.png', width=Inches(6.4))
        os.remove('plt_1.png')

        document.save('Отчет.docx')
        wdFormatPDF = 17
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(f'{self.directory}\\Отчет.docx')
        doc.SaveAs(f'{self.directory}\\Отчет.pdf', FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()
        os.remove('Отчет.docx')

    def work_with_data(self, res):
        cur_table = pd.DataFrame(pd.concat(res[0]).groupby('Должность').mean())
        if not os.path.isfile('general_table.xlsx'):
            general_table = pd.DataFrame.from_dict({"Должность": cur_table.index})
            self.remake_workers(cur_table)
            general_table.to_excel('general_table.xlsx')
        data = self.available_workers
        to_delete = set(list(cur_table.index)).difference(set(data['Должность']))
        cur_table.drop(to_delete, inplace=True)
        to_add = set(data['Должность']).difference(set(list(cur_table.index)))
        difficult = res[1]
        data_to_add = {key: -1 for key in to_add}
        temp = pd.DataFrame.from_dict(data_to_add, orient='index', columns=[f'{difficult}'])
        cur_table = pd.concat([cur_table, temp])

        print(cur_table)

        if self.active_regimes:
            if 'Вычиление оценок' in self.active_regimes:
                self.grade_calculation(cur_table, res)
                self.save_settings()
                self.prepare_settings()
            if 'Оценка результатов' in self.active_regimes and 'Вычиление оценок' in self.active_regimes:
                self.grade_estimate(cur_table, res)
            elif all([os.path.exists(self.directory + '\\general_table.xlsx'),
                      self.data_for_calculation['Already_calc'],
                      'Оценка результатов' in self.active_regimes]):
                self.grade_estimate(cur_table, res)
            elif 'Вычиление оценок' not in self.active_regimes:
                self.log_message('Оценки для сотрудников еще не рассчитаны.\nВыберете режим "Вычиление оценок"')
        else:
            self.log_message('Выберете режим работы')

        self.save_settings()
        self.prepare_settings()
        self.should_update = True
        self.should_extend = False

    def grade_calculation(self, cur_table, res):
        print(self.should_extend)
        if self.should_extend:
            general_table = pd.read_excel('general_table.xlsx')
            cur_table['Должность'] = list(cur_table.index)
            general_table = pd.merge(general_table, cur_table, how='inner', on='Должность')
            general_table.drop('Unnamed: 0', axis=1, inplace=True)
            print(general_table)
            os.remove('general_table.xlsx')
            general_table.to_excel('general_table.xlsx')
            self.log_message('Новые данные добавлены')
            self.log_message(f'Результаты находятся в таблице: {self.directory}/general_table.xlsx')
            print('Новые данные добавлены')
        if self.should_update:
            general_table = pd.read_excel('general_table.xlsx')
            difficult = res[1]
            try:
                general_table[f'{difficult}'] = cur_table.values
            except ValueError:
                self.remake_workers(cur_table)
                general_table[f'{difficult}'] = cur_table.values
            general_table.drop('Unnamed: 0', axis=1, inplace=True)
            os.remove('general_table.xlsx')
            general_table.to_excel('general_table.xlsx')
            self.log_message('Данные обновлены')
            self.log_message(f'Результаты находятся в таблице: {self.directory}\\general_table.xlsx')
            print('Данные обновлены')

    def grade_estimate(self, cur_table, res):
        general_table = pd.read_excel('general_table.xlsx')

        try:
            x = Counter(general_table[f'{res[1]}'])
            y = Counter(cur_table[f'{res[1]}'])
        except KeyError:
            self.log_message('Ошибка. Для данной сложности нет оценок')
            return

        self.make_compare_plot(x, y)
        # means = self.calc_mean(cur_table)
        # print(f'Среднее значение для группы с низкой ожидаемой подготовкой - {means[0]}\n'
        #      f'Среднее значение для группы со средней ожидаемой подготовкой - {means[1]}\n'
        #      f'Среднее значение для группы с высокой ожидаемой подготовкой - {means[2]}\n')
        bad_top = self.show_top(cur_table)
        top_info_text = ""
        for val in bad_top:
            top_info_text += f'Оценка для {val[0]}: {val[1][0]}\n'
        top_info_text += '\n'

        general_info_text = f'Сложность учения: {res[1]}\n' \
                            f'Уникальных должнсотей на учении: {len(cur_table.values)}'

        self.make_grades_table(cur_table, res)
        try:
            self.make_document(general_info_text, top_info_text)
            self.log_message('Оценка произведена')
            self.log_message(f'Отчет об оценке: {self.directory}\\Отчет.pdf')
        except PermissionError:
            self.log_message('Не удалось получить доступ к файлу')
            self.log_message('Закройте отчет и попробуйте снова')

    def save_settings(self):
        with open('data/settings.json', 'w', encoding='utf8') as settings_file:
            json.dump(self.data_for_calculation, settings_file, ensure_ascii=False)


class App(CTk.CTk):
    def __init__(self):
        super().__init__()

        self.geometry("850x605")
        self.title("Fishing calculator")
        self.resizable(False, False)

        self.tab_view = MyTabView(master=self, width=850, height=555)
        self.tab_view.grid(row=0, column=0, padx=0, pady=0)


if __name__ == '__main__':
    app = App()
    app.mainloop()

#C:\Users\Сергей\Desktop\Data\

"""
Составление оценки сложности
Интуитивный интерфейс
Интерфейс оценки границ исходов учения и доверительный интервал (отклонения) выбирает сам
Добавить разибиение по вероятностям исходов
Выбирать кого показывать через настройки (300 чекбоксов), добавить поиск
Строить гистограммы, по тем, кого выбрали
Доп столбец разбиение по 6, для каждой должности pychart
Добавить информативности, настройка границ
Добавить систему оценивания для итоговых исходов
"""